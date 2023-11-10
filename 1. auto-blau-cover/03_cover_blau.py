import os
import docx
import json
from pathlib import Path
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx import Document



def obtain_style(template_path):
    template = docx.Document(template_path)
    paragraphs = template.paragraphs
    styles = []

    for paragraph in paragraphs:
        for run in paragraph.runs:
            style_dict = {
                "text": run.text,
                "bold": run.bold,
                "font_size": run.font.size.pt if run.font.size else None,
                "style": str(run.style),
                "alignment": paragraph.alignment
            }
            styles.append(style_dict)

    return styles

def write_the_world():
    doc = docx.Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(16)
    
    file = open("tags.txt")
    lines = file.readlines()
    for i in lines:
        print(i)
        paraObj1 = doc.add_paragraph()
        paraObj1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paraObj1.style = doc.styles['Normal']
        runObj1 = paraObj1.add_run(i)
        doc.add_page_break()
    doc.save('test.docx')

def create_directory_dict(path):
    '''
    Return a JSON file with the subdirectories of the path received
    '''
    directory_dict = {}

    # Obtain al the subdirectories of the path.
    subdirs = [d for d in os.listdir(path) if os.path.isdir(os.path.join(path, d))]

    for subdir in subdirs:
        subdir_dict = {}
        subdir_dict["title"] = subdir

        subdir_path = os.path.join(path, subdir)

        subdir_subdirs = [d for d in os.listdir(subdir_path) if os.path.isdir(os.path.join(subdir_path, d))]

        subdir_dict["elements"] = subdir_subdirs

        directory_dict[subdir] = subdir_dict

    with open(os.path.join(path, "directory_dict.json"), "w") as f:
        json.dump(directory_dict, f)
    return directory_dict

def create_word_blau(path):
    """
    Create Word documents based on the contents of subdirectories.
    """
    # Create a dictionary with data from subdirectories
    directory_dict = build_directory_dict(path)

    # Create the Word document and set text style
    document = create_word_document()
    populate_document_with_data(document, directory_dict, include_elements=True)
    save_word_document(document, path, "einlegeblätter_blau.docx")

    # Create another Word document for 'einlegeblätter_gelb.docx'
    document_gelb = create_word_document()
    populate_document_with_data(document_gelb, directory_dict, include_elements=False)
    save_word_document(document_gelb, path, "einlegeblätter_gelb.docx")

def build_directory_dict(path):
    """
    Build a dictionary with titles and elements from subdirectories.
    """
    directory_dict = {}
    subdirs = [d for d in os.listdir(path) if os.path.isdir(os.path.join(path, d))]
    for subdir in subdirs:
        subdir_dict = {"title": subdir}
        subdir_path = os.path.join(path, subdir)
        subdir_subdirs = [d for d in os.listdir(subdir_path) if os.path.isdir(os.path.join(subdir_path, d))]
        subdir_dict["elements"] = subdir_subdirs
        directory_dict[subdir] = subdir_dict
    return directory_dict

def create_word_document():
    """
    Create a Word document with specific font settings.
    """
    document = Document()
    font = document.styles['Normal'].font
    font.name = 'Arial'
    font.size = Pt(16)
    return document

def populate_document_with_data(document, directory_dict, include_elements=True):
    """
    Populate the Word document with data from the directory dictionary.
    """
    for dictionary in directory_dict.values():
        document.add_paragraph(dictionary["title"], style='Title')

        if include_elements:
            for element in dictionary["elements"]:
                paragraph = document.add_paragraph(element)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        document.add_page_break()

def save_word_document(document, path, filename):
    """
    Save the Word document to a specified file.
    """
    word_file = os.path.join(path, filename)
    document.save(word_file)


# create_word_blau('/path/to/directory')

# Usage Example
# directory = r".\10_Schweißnahtdokumention\10.6_Materialzertifikate"
# create_word_blau(directory)


